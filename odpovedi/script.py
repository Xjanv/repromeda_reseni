
#!/usr/bin/env python
# coding: utf-8

# In[ ]:


#Úkol 1
#Import potřebných knihoven
import pandas as pd
import numpy as np
from scipy.stats import ttest_ind, shapiro, levene, mannwhitneyu, fisher_exact
import matplotlib.pyplot as plt
from pathlib import Path

zadani_path = Path(__file__).parent / '../zadani/transfery.csv'


# Načtení dat z CSV souboru
df = pd.read_csv(zadani_path)


# In[ ]:


#1.A) 
#Dle věku matky “vek_mother”, ve věkových kategoriích viz tabulka, vytvořte tabulku úspěšnosti embryotransferu v
#procentech dle sloupce “clinical_gravidity”, kde 1 = transfer byl úspěšný a 0 = neúspěšný. 
#Prázdné hodnoty do statistik nepočítejte.

# Filtrace dat pro věk matky
df_mother = df[df['vek_mother'] != 'X']

#Převedení na číselný formát
df_mother['vek_mother'] = pd.to_numeric(df_mother['vek_mother'], errors='coerce')

# Odstranění řádků s prázdnými hodnotami v 'clinical_gravidity'
df_mother = df_mother.dropna(subset=['clinical_gravidity'])

# Vytvoření kategorického sloupce pro věkové kategorie
bins = [0, 29, 34, 39, float('inf')]
labels = ['do 29', '30-34', '35-39', '40 a výše']
df_mother['vek_mother_category'] = pd.cut(df_mother['vek_mother'], bins=bins, labels=labels, right=True)

# Vytvoření kontingenční tabulky pro úspěšnost embryotransferu
success_table_mother = pd.crosstab(df_mother['vek_mother_category'], df_mother['clinical_gravidity'], margins=True, margins_name='Dohromady')
success_table_mother['Úspěšnost embryotransferu (%)'] = (success_table_mother[1] / success_table_mother['Dohromady']) * 100

print(success_table_mother)


# In[ ]:


#1.B) 
#Určete zda-li je věk matky statisticky významný na úspěch transferu.

#Odstranění prázdných hodnot
df_mother = df_mother.dropna(subset=['vek_mother'])

# Rozdělení dat do dvou skupin
success_group_mother = df_mother[df_mother['clinical_gravidity'] == 1]['vek_mother']
failure_group_mother = df_mother[df_mother['clinical_gravidity'] == 0]['vek_mother']

# Ověření normality dat
stat_success, p_success = shapiro(success_group_mother)
stat_failure, p_failure = shapiro(failure_group_mother)

# Srovnání rozptylů
stat_levene, p_levene = levene(success_group_mother, failure_group_mother)

# Provedení t-testu pouze pokud jsou splněny předpoklady
if p_success > 0.05 and p_failure > 0.05 and p_levene > 0.05:
    # Provedení t-testu
    t_statistic, p_value = ttest_ind(success_group_mother, failure_group_mother)

    # Vyhodnocení výsledků
    print(f"T-Statistic: {t_statistic}")
    print(f"P-Value: {p_value}")

    # Výsledek testu
    if p_value < 0.05:
        print("Rozdíl v průměrném věku mezi skupinami je statisticky významný.")
    else:
        print("Rozdíl v průměrném věku mezi skupinami není statisticky významný.")
else:
    print("Předpoklady pro t-test nejsou splněny (normalita dat nebo homogenita rozptylů).")
    
print(f"Shapiro-Wilk Test (Success Group): Statistic = {stat_success}, P-Value = {p_success}")
print(f"Shapiro-Wilk Test (Failure Group): Statistic = {stat_failure}, P-Value = {p_failure}")
print(f"Levene Test: Statistic = {stat_levene}, P-Value = {p_levene}")




#Následný MannWhitney test

# Mann-Whitneyův test
statistic, p_value = mannwhitneyu(success_group_mother, failure_group_mother, alternative='two-sided')

# Výsledky testu
print(f"Mann-Whitney U test statistic: {statistic}")
print(f"P-value: {p_value}")

# Interpretace výsledků
alpha = 0.05
if p_value < alpha:
    print("Přijímáme alternativní hypotézu: Existuje statisticky významný rozdíl mezi věkem matky a úspěchem transferu.")
else:
    print("Nemáme dostatek důkazů na odmítnutí nulové hypotézy: Neexistuje statisticky významný rozdíl mezi věkem matky a úspěchem transferu.")


# In[ ]:


#1.C) - 
#Taktéž A-B proveďte i pro věk embrya “vek_embryo”. Pokud bylo embryo darované ”f_donor” = 1, takový transfer do statistiky nepočítejte.

#Vytvoření tabulky

# Filtrace dat (odstranění řádků s 'X' ve sloupci 'vek_embryo')
df_embryo = df[df['vek_embryo'] != 'X']

# Filtrace dat (pouze řádky, kde 'f_donor' není 1)
df_embryo = df_embryo[df_embryo['f_donor'] != 1]

# Převedení sloupce 'vek_embryo' na numerický formát
df_embryo['vek_embryo'] = pd.to_numeric(df_embryo['vek_embryo'], errors='coerce')

# Filtrace dat (odstranění řádků s prázdnými hodnotami v 'clinical_gravidity')
df_embryo = df_embryo.dropna(subset=['clinical_gravidity'])


# Definice kategorií a rozdělení do kategorií 'vek_embryo_category'
bins = [0, 29, 34, 39, float('inf')]
labels = ['do 29', '30-34', '35-39', '40 a výše']
df_embryo['vek_embryo_category'] = pd.cut(df_embryo['vek_embryo'], bins=bins, labels=labels, right=True)

# Vytvoření kontingenční tabulky
success_table_embryo = pd.crosstab(df_embryo['vek_embryo_category'], df_embryo['clinical_gravidity'], margins=True, margins_name='Dohromady')

# Přidání sloupce s úspěšností embryotransferu v procentech
success_table_embryo['Úspěšnost embryotransferu (%)'] = (success_table_embryo[1] / success_table_embryo['Dohromady']) * 100

# Výpis vytvořené tabulky
print(success_table_embryo)


# In[ ]:


#1.C) 
#Taktéž A-B proveďte i pro věk embrya “vek_embryo”. Pokud bylo embryo darované ”f_donor” = 1, takový transfer do statistiky nepočítejte.
from scipy.stats import ttest_ind, shapiro, levene

#Otestování
# Vytvoření nového DataFrame pouze s potřebnými hodnotami
df_filtered = df[(df['vek_embryo'] != 'x') & (~df['vek_embryo'].isna()) & (df['f_donor'] != 1)]
df_filtered = df_filtered.dropna(subset=['vek_embryo', 'clinical_gravidity'])
df_filtered['vek_embryo'] = pd.to_numeric(df_filtered['vek_embryo'], errors='coerce')
#filtered_df = df[(df['vek_embryo'] != 'x') & (~df['vek_embryo'].isna()) & (df['f_donor'] != 1) & (~df['clinical_gravidity'].isna())]


# Ověření normality dat
stat_success, p_success = shapiro(df_filtered[df_filtered['clinical_gravidity'] == 1]['vek_embryo'])
stat_failure, p_failure = shapiro(df_filtered[df_filtered['clinical_gravidity'] == 0]['vek_embryo'])

print(f"Shapiro-Wilk Test (Success Group): Statistic = {stat_success}, P-Value = {p_success}")
print(f"Shapiro-Wilk Test (Failure Group): Statistic = {stat_failure}, P-Value = {p_failure}")

# Rozdělení dat do dvou skupin
success_group = df_filtered[df_filtered['clinical_gravidity'] == 1]['vek_embryo']
failure_group = df_filtered[df_filtered['clinical_gravidity'] == 0]['vek_embryo']

# Mann-Whitneyův test
statistic, p_value = mannwhitneyu(success_group, failure_group, alternative='two-sided')

# Výsledky testu
print(f"Mann-Whitney U test statistic: {statistic}")
print(f"P-value: {p_value}")

# Interpretace výsledků
alpha = 0.05
if p_value < alpha:
    print("Přijímáme alternativní hypotézu: Existuje statisticky významný rozdíl mezi věkem embryí a úspěchem transferu.")
else:
    print("Nemáme dostatek důkazů na odmítnutí nulové hypotézy: Neexistuje statisticky významný rozdíl mezi věkem embryí a úspěchem transferu.")


# In[ ]:


#1.D) 
#Vytvořte tabulku s počty transferů dle použité genetické metody "genetic_method” viz tabulka. 

conditions = [
    (df['genetic_method'] == 'PGT-A'),
    (df['genetic_method'] == 'PGT-SR'),
    (df['genetic_method'] == 'Karyomapping'),
    (df['genetic_method'] == 'OneGene'),
    (df['genetic_method'].isna()) & (~df['clinical_gravidity'].isna()),
]
choices = ['PGT-A', 'PGT-SR', 'Karyomapping', 'OneGene', 'Bez genetické metody']
df['genetic_method_category'] = np.select(conditions, choices, default='Ostatní')

# Vytvoření kontingenční tabulky bez sloupců s hodnotami 0 a 1
result_table = pd.crosstab(df['genetic_method_category'], df['clinical_gravidity'], margins=True, margins_name='Total')
result_table = result_table.loc[:, result_table.columns.difference([0.0, 1.0])]

print(result_table)


# In[ ]:


#1.E) 
#Určete statistickou významnost pohlaví embrya “sex” – XX/XY na úspěch klinické gravidity dle sloupce “clinical_gravidity”, 
#kde 1 = transfer byl úspěšný a 0 = neúspěšný. Prázdné hodnoty do statistik nepočítejte.

# Filtrace dat
df_sex = df[(df['sex'].isin(['XX', 'XY']))  & (~df['clinical_gravidity'].isna())]

# Kontrola výskytu hodnot v kategorii 'sex'
#print(df_sex['sex'].value_counts())

# Kontrola výskytu hodnot v kategorii 'clinical_gravidity'
#print(df_sex['clinical_gravidity'].value_counts())

# Kontingenční tabulka
contingency_table = pd.crosstab(df_sex['sex'], df_sex['clinical_gravidity'])

# Fisherův exaktní test
odds_ratio, p = fisher_exact(contingency_table)

# Výsledky testu
print(f"Odds Ratio: {odds_ratio}")
print(f"P-value: {p}")

# Interpretace výsledků
alpha = 0.05
if p < alpha:
    print("Přijímáme alternativní hypotézu: Existuje statisticky významný vztah mezi pohlavím embrya a úspěchem klinické gravidity.")
else:
    print("Nemáme dostatek důkazů na odmítnutí nulové hypotézy: Neexistuje statisticky významný vztah mezi pohlavím embrya a úspěchem klinické gravidity.")


# In[128]:


#1.F) 
#Z výsledných tabulek z úkolu A a D vytvořte a uložte grafy ve formátu .png, kde na ose x bude první a na ose y druhý řádek tabulky.



# Vytvoření sloupcového grafu z úlohy D
ax = result_table['Total'][:-1].plot(kind='bar', color='darkgray', edgecolor='black')

# Přidání popisků
plt.title('Počet transferů podle genetické metody')
plt.xlabel('Genetická metoda')
plt.ylabel('Celkový počet transferů')

# Nastavení limitu na ose Y
plt.ylim(0, 1000)

# Uložení grafu do souboru .png na plochu
odpovedi_path = Path(__file__).parent / 'sloupcovy_graf1.png'
plt.savefig(odpovedi_path, bbox_inches='tight')
plt.show()


# # Vytvoření sloupcového grafu z úlohy A
plt.bar(success_table_mother.index, success_table_mother['Úspěšnost embryotransferu (%)'])

# Nastavení popisků os
plt.xlabel('Věk matky')
plt.ylabel('Úspěšnost embryotransferu (%)')
plt.title('Úspěšnost embryotransferu podle věku matky')

# Nastavení limitu na ose Y
plt.ylim(0, 100)

# Uložení grafu do souboru .png na plochu
odpovedi_path = Path(__file__).parent / 'sloupcovy_graf2.png'
plt.savefig(odpovedi_path, bbox_inches='tight')

# Zobrazení grafu
plt.show()


# In[129]:


# Úkol 2
#2.	Vytvořte script, který bude vytvářet jednoduchý .docx dokument obsahující nadpis, tučným a zarovnaný 
#na střed “Výsledný protokol genetického vyšetření”, a dále pak tabulku, která se vyplní dle tří vstupních argumentů 

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_genetic_protocol(doc, name, birth_number, sampling_date):
    # Nadpis
    title = doc.add_heading('Výsledný protokol genetického vyšetření', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Nový odstavec
    doc.add_paragraph()

    # Tabulka
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Pt(150)
    table.columns[1].width = Pt(250)

    # Nastavení stylu pro buňky tabulky
    for row in table.rows:
        for cell in row.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    # Nastavení obsahu tabulky
    table.cell(0, 0).text = 'Jméno a příjmení:'
    table.cell(0, 1).text = name
    table.cell(1, 0).text = 'Rodné číslo:'
    table.cell(1, 1).text = birth_number
    table.cell(2, 0).text = 'Datum odběru:'
    table.cell(2, 1).text = sampling_date

# Příklad použití
odpovedi_path = Path(__file__).parent / 'vysledny_protokol.docx'

doc = Document()
create_genetic_protocol(doc, 'Jan Novák', '123456/7890', '2024-02-10')
doc.save(odpovedi_path)


# In[ ]:




